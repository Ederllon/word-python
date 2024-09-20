
from docx import Document

# Criar um novo documento
doc = Document()
doc.add_heading('Título do Documento', level=1)
doc.add_paragraph('Este é um parágrafo no documento.')
doc.save('documento.docx')
